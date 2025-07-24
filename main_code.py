from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


# Constants
constants = [
    ("Closing speed V", "3000 ft/s (≈914 m/s)"),
    ("Time constant τ", "2 s"),
    ("Process noise intensity W", "100 ( m²/s³ )"),
    ("Meas. noise coeff. R₁", "15×10⁻⁶ rad²/s"),
    ("Meas. noise coeff. R₂", "1.67×10⁻³ rad² s"),
    ("Control‑effort weight b", "1.52×10⁻²"),
    ("P₂₂(0)", "16 (m/s)²"),
    ("P₃₃(0)", "400 (m/s²)²"),
    ("Final time t_f", "10 s")
]

# Create new document
doc = Document()


# ---------- 1 Introduction ----------
doc.add_heading('1   Introduction', level=1)
intro = (
    "Missile–target interception is a classical real‑time decision problem: the pursuer must steer "
    "so that the line‑of‑sight (LOS) separation y(t) crosses zero at the terminal time t_f while its "
    "radial closure speed V remains essentially constant. In practice, the target executes evasive "
    "turns and the seeker’s angle measurement is noisy; hence a purely deterministic guidance law is "
    "insufficient. A stochastic optimal‑control framework—Linear Quadratic Gaussian (LQG)—allows us "
    "to design a feedback law that explicitly accounts for process noise (unknown target manoeuvres) "
    "and measurement noise while keeping control effort within realistic bounds.\n\n"
    "**Goal of the project.** Design, simulate, and analyse a continuous‑time LQG controller that "
    "minimises the *expected* terminal miss distance ½ y(t_f)² and the cumulative control energy "
    "½ b∫₀^{t_f} u(t)²dt. The task synthesises all core topics of the course—continuous‑time LQR, "
    "Kalman filtering, and the separation principle—into a single capstone problem."
)
doc.add_paragraph(intro)

# ---------- 2 Problem Statement ----------
doc.add_heading('2   Problem Statement and Requirements', level=1)
ps = (
    "We control a pursuer missile whose lateral acceleration u(t)=a_P(t) must null the LOS displacement "
    "y(t) at a fixed final time t_f=10 s. The relative kinematics are represented by a three‑state model:\n"
    "• y — lateral displacement (ft or m)\n"
    "• v — relative lateral velocity (ft/s or m/s)\n"
    "• a_T — target lateral acceleration (ft/s² or m/s²), modelled as a first‑order Gauss–Markov process.\n\n"
    "Only the noisy LOS angle θ(t)=y(t)/(V(t_f−t)) is measured. Target acceleration and velocity are "
    "unobserved. The designer must:\n"
    "1. Formulate the system in continuous‑time state‑space form with process and measurement noise.\n"
    "2. Derive the optimal continuous‑time Kalman filter to estimate (y, v, a_T).\n"
    "3. Derive the optimal continuous‑time LQR controller given the terminal‑plus‑integral quadratic cost.\n"
    "4. Combine them via the separation principle, simulate ≥1000 Monte‑Carlo runs, and report statistics "
    "on terminal miss and control effort.\n\n"
    "Assumptions: constant closing speed V, small LOS angle so sin θ≈θ, and no actuator saturation."
)
doc.add_paragraph(ps)

doc.add_heading('Numerical Constants', level=2)
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Symbol / Description'
hdr_cells[1].text = 'Value'
for name, val in constants:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = val

# ---------- 3 Mathematical Modelling ----------
doc.add_heading('3   Mathematical Modelling', level=1)

doc.add_heading('3.1 Continuous‑Time State‑Space Representation', level=2)
eq1 = (
    "State vector   x(t) = [ y, v, a_T ]ᵀ\n"
    "Control        u(t) = a_P(t)\n\n"
    "ẋ = A x + B u + G w,  z = H(t) x + m,\n"
    "with\n"
    "A = [[0, 1, 0],\n"
    "     [0, 0, 1],\n"
    "     [0, 0, −1/τ]],\n"
    "B = [0, 1, 0]ᵀ, G = [0, 0, 1]ᵀ √W,\n"
    "H(t) = [1/(V(t_f−t)), 0, 0]."
)
doc.add_paragraph(eq1)

doc.add_heading('3.2 Noise Modelling', level=2)
noise_text = (
    "Process noise w(t) is scalar, zero‑mean white with intensity W (ft/s²)² s or (m/s²)² s. It drives the "
    "target acceleration state a_T, capturing unpredictable manoeuvres.\n\n"
    "Measurement noise m(t) is zero‑mean white with *time‑varying* variance\n"
    "  Rₘ(t) = R₁ + R₂/(t_f−t)²,\n"
    "which grows as the engagement nears completion, reflecting seeker angle‑to‑LOS conversion uncertainty."
)
doc.add_paragraph(noise_text)

doc.add_heading('3.3 Performance Index', level=2)
eq2 = (
    "J = 𝔼[ ½ y(t_f)² + ½ b ∫₀^{t_f} u(t)² dt ].\n"
    "The first term penalises terminal miss distance; the second penalises control energy with weight b."
)
doc.add_paragraph(eq2)

# ---------- 4 Theoretical Background ----------
doc.add_heading('4   Theoretical Background', level=1)

doc.add_heading('4.1 Linear Quadratic Regulator (LQR)', level=2)
lqr_text = (
    "For a deterministic linear system ẋ = A x + B u and quadratic cost "
    "J = ½ x(t_f)ᵀ S_f x(t_f) + ½ ∫₀^{t_f} (xᵀQx + uᵀRu)dt, the optimal control is "
    "u = −K(t) x where K(t) = R⁻¹BᵀS(t). S(t) solves the *control Riccati ODE*\n"
    "  −Ṡ = AᵀS + S A − S B R⁻¹ Bᵀ S + Q, S(t_f) = S_f."
)
doc.add_paragraph(lqr_text)

doc.add_heading('4.2 Kalman Filtering', level=2)
kf_text = (
    "For ẋ = A x + G w, z = H x + m with cov(w)=W δ and cov(m)=Rₘ(t) δ, the minimum‑variance "
    "estimate x̂ obeys\n"
    "  ẋ̂ = A x̂ + B u + L(t)(z − H x̂),\n"
    "where the *filter gain* L(t) = P(t) Hᵀ Rₘ⁻¹ and P(t) is the estimation error covariance solving\n"
    "  Ṗ = A P + P Aᵀ − P Hᵀ Rₘ⁻¹ H P + G W Gᵀ, P(0) = diag(0, P₂₂(0), P₃₃(0))."
)
doc.add_paragraph(kf_text)

doc.add_heading('4.3 Separation Principle & LQG', level=2)
sep_text = (
    "Because the plant, cost, and noise are all linear‑quadratic‑Gaussian, the optimal controller under "
    "partial information is obtained by combining the LQR gain K(t) (designed for the full‑state case) "
    "with the Kalman estimate x̂(t). The closed‑loop policy\n"
    "  u(t) = −K(t) x̂(t)\n"
    "is optimal and stabilising provided (A,B) is stabilisable and (A,H) is detectable—both satisfied here."
)
doc.add_paragraph(sep_text)

# ---------- 5 Solution Methodology ----------
doc.add_heading('5   Solution Methodology', level=1)

# ---------- 5.1 Control Riccati ODE ----------
doc.add_heading('5.1 Control Riccati ODE', level=2)
text_51 = (
    "In the Linear Quadratic Regulator (LQR) framework, we assume full state information x(t) is available. "
    "The goal is to find a time-varying feedback gain K(t) such that the control law\n"
    "  u(t) = −K(t) x(t)\n"
    "minimises the performance index\n"
    "  J = 𝔼[½ y(t_f)² + ½ b ∫₀^{t_f} u(t)² dt].\n\n"
    "We express this as a quadratic cost with terminal weight\n"
    "  S_f = diag(1, 0, 0),  R = b = 1.52×10⁻².\n\n"
    "The optimal gain is derived from the control Riccati differential equation:\n"
    "  −Ṡ(t) = Aᵀ S + S A − S B R⁻¹ Bᵀ S,\n"
    "with terminal condition S(t_f) = S_f.\n\n"
    "System matrices used:\n"
    "  A = [[0, 1, 0], [0, 0, 1], [0, 0, −1/τ]]\n"
    "  B = [0, 1, 0]ᵀ\n"
    "  R = b = 1.52 × 10⁻²\n\n"
    "This ODE is integrated backward in time from t = t_f to 0. In Python, we use scipy.integrate.solve_ivp "
    "with t_span = [t_f, 0]. Matrix S is flattened into a vector during integration and reshaped at each step."
)
doc.add_paragraph(text_51)

# ---------- 5.2 Filter Riccati ODE ----------
doc.add_heading('5.2 Filter Riccati ODE', level=2)
text_52 = (
    "In the Kalman filtering framework, we estimate the full state x(t) = [y, v, a_T]ᵀ based on noisy measurements "
    "of the LOS angle θ(t). The estimator has the form:\n"
    "  ẋ̂(t) = A x̂ + B u + L(t)(z − H x̂),\n"
    "where L(t) is the time-varying Kalman gain computed from the error covariance matrix P(t).\n\n"
    "The error covariance P(t) evolves forward in time according to the filter Riccati differential equation:\n"
    "  Ṗ(t) = A P + P Aᵀ − P Hᵀ Rₘ⁻¹ H P + G W Gᵀ,\n"
    "with initial condition:\n"
    "  P(0) = diag(0, P₂₂(0), P₃₃(0)).\n\n"
    "Measurement noise is time-dependent:\n"
    "  Rₘ(t) = R₁ + R₂ / (t_f − t)²,\n"
    "which reflects increasing angular noise as the missile approaches the target.\n\n"
    "System matrices:\n"
    "  A = [[0, 1, 0], [0, 0, 1], [0, 0, −1/τ]]\n"
    "  G = [0, 0, 1]ᵀ\n"
    "  H(t) = [1 / (V(t_f − t)), 0, 0]\n"
    "  W = 100,  R₁ = 15×10⁻⁶,  R₂ = 1.67×10⁻³\n\n"
    "As with the control Riccati ODE, this equation is integrated using solve_ivp, "
    "but in the forward direction from t = 0 to t = t_f. The matrix P is flattened into a vector "
    "for integration and reshaped during evaluation."
)
doc.add_paragraph(text_52)


# ---------- 5.3 Numerical Integration Strategy ----------
doc.add_heading('5.3 Numerical Integration Strategy', level=2)
text_53 = (
    "The Riccati differential equations governing both the control matrix S(t) and the estimation error P(t) "
    "are time-dependent matrix ODEs that require careful numerical integration.\n\n"
    "• The control Riccati equation for S(t) is integrated *backward in time* from t = t_f to t = 0. "
    "This requires reversing the time axis in solve_ivp using t_span = [t_f, 0]. The terminal condition is "
    "S(t_f) = diag(1, 0, 0). At each step, the matrix S is vectorised (flattened) before integration and reshaped afterward.\n\n"
    "• The filter Riccati equation for P(t) is integrated *forward in time* from t = 0 to t = t_f using "
    "t_span = [0, t_f]. The initial condition is P(0) = diag(0, P₂₂(0), P₃₃(0)). The measurement noise matrix Rₘ(t) "
    "is evaluated at each time step based on the expression:\n"
    "  Rₘ(t) = R₁ + R₂ / (t_f − t)².\n\n"
    "We recommend using SciPy’s solve_ivp with an adaptive time-stepping method like 'RK45' or 'LSODA', "
    "with absolute and relative tolerances set to 1e‑8 or tighter, to ensure numerical stability and accuracy.\n\n"
    "In both cases, the right-hand side functions of the ODEs are defined using flattened versions of the symmetric "
    "matrices S and P. At the end of integration, the results are reshaped into 3×3 matrix trajectories. "
    "These matrices are then used to compute the control gain K(t) and the Kalman gain L(t) at each instant."
)
doc.add_paragraph(text_53)


# ---------- 6. Implementation in Python ----------
doc.add_heading('6. Implementation in Python', level=1)

doc.add_paragraph(
    'All numerical simulations and controller evaluations were carried out using Python 3.x, '
    'leveraging scientific libraries such as NumPy, SciPy, and Matplotlib. The code was implemented '
    'in a modular fashion to ensure clarity, maintainability, and extensibility.'
)

doc.add_heading('6.1 Core Functions', level=2)

doc.add_paragraph(
    'The core dynamics and noise models were encapsulated in a set of foundational functions. '
    'The function state_matrices() defines the system matrices A, B, and G based on the known structure '
    'of the target–pursuer model. The time-dependent measurement noise covariance is modeled by the function Rm(t), '
    'while the time-varying observation matrix H(t) is computed via the helper function H(t), both designed to '
    'prevent singularities by clamping the time-to-go variable near zero.'
)

doc.add_paragraph(
    'The right-hand side (RHS) of the Riccati differential equation for the Kalman filter is implemented in '
    'filter_riccati_rhs(t, P_flat), which propagates the estimation covariance matrix P(t) forward in time. '
    'In parallel, the RHS for the LQR optimal control Riccati equation is implemented in control_riccati_rhs(t, S_flat), '
    'propagating the cost-to-go matrix S(t) backward in time. These Riccati equations are solved numerically using '
    'SciPy’s solve_ivp() integrator with high accuracy tolerances.'
)

doc.add_paragraph(
    'The function simulate_closed_loop(seed, time_grid, S_traj, P_traj) executes one Monte Carlo realization of the '
    'closed-loop system using the Euler–Maruyama method. It simulates the dynamics of the system, applies the LQG controller '
    'in real time based on the state estimate, and accumulates both the terminal miss distance y(tf) and the cumulative '
    'quadratic cost J. Random seeds are used to introduce stochastic variability in initial conditions and process noise, '
    'ensuring statistical robustness.'
)

doc.add_heading('6.2 Monte Carlo Driver', level=2)

doc.add_paragraph(
    'A dedicated driver function, run_monte_carlo(N), executes N = 1000 independent realizations of the closed-loop system. '
    'It first integrates the Riccati equations to obtain trajectories P(t) and S(t), then loops over N trials, each with '
    'randomized initial conditions and disturbances. The outputs collected are the final lateral miss distances y(tf) and '
    'the cumulative costs J. These results are then visualized through histograms, and key statistics such as mean and '
    'standard deviation are printed for both metrics.'
)
doc.add_paragraph(
    'In addition to scalar summaries, the simulation framework also computes full time histories of the state and control '
    'trajectories for each realization. This allows the post-processing functions to generate additional performance diagnostics, '
    'such as the mean trajectory and standard deviation envelope of y(t) over time (Mean ± 1σ), and a representative control signal '
    'a_P(t) from a single trial. These enrich the interpretation of the Monte Carlo outcomes beyond terminal statistics.'
)

doc.add_heading('6.3 Validation Checks', level=2)

doc.add_paragraph(
    'To ensure numerical and physical validity, a set of validation checks were performed. First, the positive '
    'semi-definiteness of the Riccati matrices P(t) and S(t) was verified for all time steps by inspecting their eigenvalues. '
    'This confirms that the estimation error covariance and the cost-to-go function remain physically meaningful throughout '
    'the simulation. Second, convergence and stability of the integration routines were ensured by using tight absolute and '
    'relative tolerances in solve_ivp(). The entire pipeline was tested for consistency, reproducibility, and correctness.'
)

# ---------- 7. Simulation Results ----------
doc.add_heading('7. Simulation Results', level=1)

doc.add_paragraph(
    'This section presents the outcomes of the Monte Carlo simulations used to evaluate the performance of the LQG controller. '
    'Each realization simulated the stochastic dynamics of the guidance system under measurement and process noise, and applied the '
    'estimated feedback control law in real time. A total of 1000 trials were conducted to ensure statistical robustness.'
)

doc.add_paragraph(
    'The primary metrics of interest are the terminal lateral miss distance y(tf) and the total quadratic cost J. '
    'Figure 1 displays the evolution of the lateral position y(t), averaged across all realizations, along with the ±1 standard deviation envelope. '
    'This provides insight into how tightly the closed-loop system converges toward the desired terminal condition over time.'
)

doc.add_paragraph(
    'Figure 2 shows a representative control input a_P(t) taken from a single trial. This illustrates the structure and magnitude of the control actions '
    'applied by the LQG strategy to correct for system uncertainties and disturbances throughout the engagement.'
)

doc.add_paragraph(
    'Figure 3 displays a histogram of the terminal miss distances y(tf) across all 1000 simulations. The distribution is centered near zero, indicating that the '
    'controller achieves its guidance objective on average. However, the spread reflects the influence of initial condition uncertainties and noise.'
)

doc.add_paragraph(
    'Figure 4 presents the histogram of the total cost J incurred per trial. The right-skewed shape of the distribution is consistent with the fact that larger control '
    'efforts are required in more adverse scenarios, which occur with lower probability.'
)

doc.add_paragraph(
    'Overall, the simulation results confirm that the LQG controller performs effectively in this stochastic guidance context. '
    'It achieves accurate terminal guidance while maintaining reasonable control effort, and remains robust across a wide range of disturbances.'
)



doc.add_paragraph(
    'This section presents the results obtained from the Monte Carlo simulations conducted over '
    '1,000 independent trials. The figures below illustrate key performance metrics of the closed-loop system under stochastic conditions.'
)

# Figure 1
doc.add_picture('figure1.png', width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Figure 1: Time evolution of the mean lateral position y(t) with shaded area representing ±1 standard deviation.',
                  style='Caption')

# Figure 2
doc.add_picture('figure2.png', width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Figure 2: Representative control effort history a_P(t) from a single Monte Carlo realization.',
                  style='Caption')

# Figure 3
doc.add_picture('figure3.png', width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Figure 3: Histogram of final miss distances y(tf) across all Monte Carlo trials.',
                  style='Caption')

# Figure 4
doc.add_picture('figure4.png', width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Figure 4: Histogram of total cost J, reflecting variability in control performance across trials.',
                  style='Caption')


# Save the document
path = '/home/maayan/Maayan/OptimalControl/missile_lqg_project_sections_filled.docx'
doc.save(path)

path
